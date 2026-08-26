from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "resume_original.docx"
OUTPUT = ROOT / "简历_美化版.docx"

NAVY = "17324D"
TEAL = "137C75"
INK = "25323D"
MUTED = "667681"
LIGHT = "EDF3F5"
RULE = "CAD6DB"
WHITE = "FFFFFF"
FONT = "Microsoft YaHei"
CONTENT_DXA = 10200


def set_font(run, size=8.8, bold=False, color=INK, italic=False):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def set_cell_margins(cell, top=50, start=80, bottom=50, end=80):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Cm(widths[idx] / 1440 * 2.54)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def remove_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        borders.append(el)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.first_child_found_in("w:shd")
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def paragraph_border(paragraph, edge, color, size=8, space=3):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    border = OxmlElement(f"w:{edge}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), str(size))
    border.set(qn("w:space"), str(space))
    border.set(qn("w:color"), color)
    borders.append(border)


def add_text(paragraph, text, size=8.8, bold=False, color=INK, italic=False):
    run = paragraph.add_run(text)
    set_font(run, size, bold, color, italic)
    return run


def keep_with_next(paragraph):
    paragraph._p.get_or_add_pPr().append(OxmlElement("w:keepNext"))


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def add_section(doc, title, page_break=False):
    p = doc.add_paragraph(style="Heading 1")
    if page_break:
        p.paragraph_format.page_break_before = True
    p.paragraph_format.left_indent = Pt(6)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3.5)
    p.paragraph_format.keep_with_next = True
    add_text(p, title, 11.2, True, NAVY)
    paragraph_border(p, "left", TEAL, size=18, space=4)
    paragraph_border(p, "bottom", RULE, size=5, space=3)
    return p


def add_body(doc, text, after=2.0, bold=False, color=INK, keep=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.12
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.keep_together = True
    if keep:
        p.paragraph_format.keep_with_next = True
    add_text(p, text, 8.7, bold, color)
    return p


def add_bullet(doc, text, size=8.45, after=1.5):
    clean = re.sub(r"^[•·\s]+", "", text)
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.48)
    p.paragraph_format.first_line_indent = Cm(-0.24)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.08
    p.paragraph_format.keep_together = True
    add_text(p, clean, size, False, INK)
    return p


def add_label_value(doc, label, value, after=2.0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.1
    p.paragraph_format.keep_together = True
    add_text(p, label, 8.7, True, TEAL)
    if value:
        add_text(p, value, 8.7, False, INK)
    return p


def add_project_title(doc, text):
    date_match = re.search(r"(20\d{2}\.\d{2}\s*[–-]\s*(?:20\d{2}\.\d{2}|至今))\s*$", text)
    date_text = date_match.group(1) if date_match else ""
    title_text = text[: date_match.start()].strip() if date_match else text.strip()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.tab_stops.add_tab_stop(Cm(18.0), WD_TAB_ALIGNMENT.RIGHT)
    add_text(p, title_text, 9.5, True, NAVY)
    if date_text:
        add_text(p, "\t" + date_text, 8.4, True, TEAL)
    return p


def split_label(text):
    match = re.match(r"^([^:：]+[:：])(.*)$", text)
    return (match.group(1), match.group(2).strip()) if match else ("", text)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_text(paragraph, "吴志鸿  |  Java / Agent 应用开发  |  ", 7.5, False, MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:eastAsia"), FONT)
    fonts.set(qn("w:ascii"), FONT)
    fonts.set(qn("w:hAnsi"), FONT)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), MUTED)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "15")
    r_pr.extend([fonts, color, size])
    run.append(r_pr)
    text = OxmlElement("w:t")
    text.text = "1"
    run.append(text)
    fld.append(run)
    paragraph._p.append(fld)


def build():
    src = Document(SOURCE)
    text = {i: p.text.strip() for i, p in enumerate(src.paragraphs)}

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.25)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.header_distance = Cm(0.55)
    section.footer_distance = Cm(0.55)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(8.8)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.line_spacing = 1.12
    for style_name in ("Heading 1", "Heading 2", "Heading 3", "List Bullet"):
        style = styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)

    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.paragraph_format.space_before = Pt(0)
    footer_p.paragraph_format.space_after = Pt(0)
    paragraph_border(footer_p, "top", RULE, size=4, space=3)
    add_page_number(footer_p)

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(1)
    add_text(kicker, "个人简历  /  SOFTWARE ENGINEERING", 7.7, True, TEAL)

    header = doc.add_table(rows=1, cols=2)
    set_table_geometry(header, [3700, 6500])
    remove_table_borders(header)
    for cell in header.rows[0].cells:
        set_cell_margins(cell, top=0, bottom=30, start=0, end=0)
    p = header.cell(0, 0).paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    add_text(p, "吴志鸿", 21.5, True, NAVY)
    p = header.cell(0, 1).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    add_text(p, "Java工程师  /  Agent应用开发工程师", 10.1, True, TEAL)

    contact = doc.add_table(rows=3, cols=4)
    set_table_geometry(contact, [760, 3900, 760, 4780])
    remove_table_borders(contact)
    labels_values = [
        ("姓名", "吴志鸿", "籍贯", "贵州省 遵义市"),
        ("年龄", "23岁", "民族", "汉族"),
        ("手机", "18586746339", "邮箱", "wzhy123zz@163.com"),
    ]
    for row, values in zip(contact.rows, labels_values):
        for idx, value in enumerate(values):
            cell = row.cells[idx]
            shade_cell(cell, LIGHT)
            set_cell_margins(cell, top=45, bottom=45, start=80, end=80)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_text(p, value, 8.15 if idx % 2 == 0 else 8.5, idx % 2 == 0, TEAL if idx % 2 == 0 else INK)

    add_section(doc, text[7])
    edu = doc.add_table(rows=1, cols=3)
    set_table_geometry(edu, [2200, 4300, 3700])
    remove_table_borders(edu)
    values = ["2023.09-2027.06", "沈阳工业大学（一本）", "电子与计算机工程专业"]
    for idx, value in enumerate(values):
        cell = edu.cell(0, idx)
        set_cell_margins(cell, top=0, bottom=20, start=0, end=40)
        p = cell.paragraphs[0]
        if idx == 2:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_text(p, value, 8.8, idx == 1, NAVY if idx == 1 else INK)
    add_label_value(doc, "GPA：", text[9].split(":", 1)[1].strip())
    courses = text[10] + text[11]
    add_label_value(doc, "主修课程：", courses.split("：", 1)[1].strip(), after=1.5)

    add_section(doc, text[12])
    for idx in range(14, 23):
        add_bullet(doc, text[idx], size=8.35, after=1.2)

    add_section(doc, text[23])
    for idx in range(25, 36):
        add_bullet(doc, text[idx], size=8.25, after=1.15)

    add_section(doc, text[37], page_break=True)
    add_project_title(doc, text[38])
    label, value = split_label(text[39])
    add_label_value(doc, label, value, after=1.4)
    label, value = split_label(text[40])
    add_label_value(doc, label, value, after=1.6)
    add_label_value(doc, "项目简介：", text[42], after=2.0)
    add_label_value(doc, text[43], "", after=1.0)
    for idx in range(44, 49):
        add_bullet(doc, text[idx], size=8.35, after=1.35)

    add_section(doc, text[50])
    add_project_title(doc, text[51])
    for idx in (52,):
        label, value = split_label(text[idx])
        add_label_value(doc, label, value, after=1.4)
    label, value = split_label(text[53] + text[54])
    add_label_value(doc, label, value, after=1.6)
    add_label_value(doc, "项目简介：", text[56], after=2.0)
    add_label_value(doc, text[57], "", after=1.0)
    for idx in range(58, 63):
        add_bullet(doc, text[idx], size=8.35, after=1.35)

    add_section(doc, text[63])
    for idx in range(65, 68):
        add_bullet(doc, text[idx], size=8.35, after=1.25)

    core = doc.core_properties
    core.title = "吴志鸿 - Java工程师 / Agent应用开发工程师简历"
    core.subject = "个人简历"
    core.author = "吴志鸿"
    core.keywords = "Java, Spring Cloud, Agent, RAG, 微服务"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
